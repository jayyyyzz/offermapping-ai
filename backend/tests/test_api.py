from __future__ import annotations

import io
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from fastapi.testclient import TestClient

import backend.app as app_module


DISABLED_MODEL = {
    "base_url": "",
    "api_key": "",
    "model": "",
    "profile": "",
}


class FakeHotspotResponse:
    def __init__(self, payload: dict) -> None:
        self.payload = payload

    def raise_for_status(self) -> None:
        return None

    def json(self) -> dict:
        return self.payload


def hotspot_payload(original_url: str = "https://example.com/story") -> dict:
    return {
        "items": [
            {
                "id": "hotspot-1",
                "title": "A traceable AI release",
                "summary": "A short verified summary.",
                "source": {"name": "Example News"},
                "links": {
                    "original": original_url,
                    "aihot": "https://aihot.virxact.com/items/hotspot-1",
                },
                "publishedAt": "2026-08-05T08:25:11.000Z",
                "category": "ai-models",
                "score": 79,
            }
        ]
    }


def insight_payload(image_url: str = "https://images.example.com/paper.png") -> dict:
    return {
        "items": [
            {
                "id": "https://www.ai-insight.org/news/14469",
                "url": "https://arxiv.org/abs/2608.03457",
                "external_url": "https://arxiv.org/abs/2608.03457",
                "title": "A traceable diffusion paper",
                "summary": "A summary from the public JSON feed.",
                "date_published": "2026-08-05T09:00:00.000Z",
                "tags": ["研究"],
                "authors": [{"name": "HuggingFace Daily Papers"}],
                "image": image_url,
            }
        ]
    }


class ApiSecurityTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.original_db_path = app_module.DB_PATH
        app_module.DB_PATH = Path(self.temp_dir.name) / "test.db"
        app_module.init_db()
        app_module.rate_limiter.reset()
        app_module.analysis_budget.reset()
        self.client = TestClient(app_module.app)

    def tearDown(self) -> None:
        self.client.close()
        app_module.DB_PATH = self.original_db_path
        self.temp_dir.cleanup()

    def register(self, email: str) -> str:
        response = self.client.post(
            "/api/auth/register",
            json={"email": email, "password": "correct-horse-123"},
        )
        self.assertEqual(response.status_code, 200)
        return response.json()["token"]

    def test_analysis_is_private_to_its_owner(self) -> None:
        token_a = self.register("owner-a@example.com")
        token_b = self.register("owner-b@example.com")
        headers_a = {"Authorization": f"Bearer {token_a}"}
        headers_b = {"Authorization": f"Bearer {token_b}"}
        jd = (
            "AI 应用开发工程师，负责使用 Python、FastAPI 和 Docker 构建大模型应用，"
            "要求具备评测、错误分析、部署与可观测性经验。"
        )
        resume = (
            "计算机专业，使用 Python 和 FastAPI 开发过检索增强问答项目，完成 Docker "
            "容器化、单元测试与失败案例分析，并维护项目文档和 GitHub 仓库。"
        ) * 2

        with patch.object(app_module, "model_config", return_value=DISABLED_MODEL):
            created = self.client.post(
                "/api/analyses",
                headers=headers_a,
                json={"jd": jd, "resume": resume},
            )

        self.assertEqual(created.status_code, 200)
        analysis_id = created.json()["analysisId"]
        self.assertEqual(
            self.client.get(f"/api/analyses/{analysis_id}", headers=headers_b).status_code,
            404,
        )
        self.assertEqual(self.client.get("/api/analyses", headers=headers_b).json(), [])
        self.assertEqual(
            self.client.get(f"/api/analyses/{analysis_id}", headers=headers_a).status_code,
            200,
        )

    def test_tampered_token_is_rejected(self) -> None:
        token = self.register("token-owner@example.com")
        tampered = f"{token[:-1]}{'a' if token[-1] != 'a' else 'b'}"
        response = self.client.get(
            "/api/analyses",
            headers={"Authorization": f"Bearer {tampered}"},
        )
        self.assertEqual(response.status_code, 401)

    def test_request_id_and_feedback_event_endpoints(self) -> None:
        token = self.register("telemetry@example.com")
        headers = {"Authorization": f"Bearer {token}", "X-Request-ID": "test-request-123"}
        event = self.client.post(
            "/api/events",
            headers=headers,
            json={"event": "analysis.opened", "metadata": {"surface": "history"}},
        )
        self.assertEqual(event.status_code, 200)
        self.assertEqual(event.json()["requestId"], "test-request-123")
        self.assertEqual(event.headers["X-Request-ID"], "test-request-123")

        feedback = self.client.post(
            "/api/feedback",
            headers=headers,
            json={"rating": "up", "comment": "Useful"},
        )
        self.assertEqual(feedback.status_code, 200)
        self.assertTrue(feedback.json()["ok"])

    def test_analysis_job_polling_and_account_deletion(self) -> None:
        token = self.register("job-owner@example.com")
        headers = {"Authorization": f"Bearer {token}"}
        jd = "AI 应用开发工程师，负责 Python、FastAPI、RAG、评测与 Docker 部署。" * 2
        resume = "使用 Python 和 FastAPI 完成过问答项目，记录评测结果并维护 GitHub 仓库。" * 3

        async def fake_analysis(_jd: str, _resume: str):
            return ({"role": "AI 应用开发工程师", "score": 60, "source": "rules"}, [])

        with patch.object(app_module, "perform_analysis", side_effect=fake_analysis):
            created = self.client.post("/api/analysis-jobs", headers=headers, json={"jd": jd, "resume": resume})
        self.assertEqual(created.status_code, 200)
        job_id = created.json()["jobId"]
        status = self.client.get(f"/api/analysis-jobs/{job_id}", headers=headers)
        self.assertEqual(status.status_code, 200)
        self.assertEqual(status.json()["status"], "completed")
        self.assertIsNotNone(status.json()["result"]["analysisId"])

        deleted = self.client.delete("/api/account", headers=headers)
        self.assertEqual(deleted.status_code, 200)
        self.assertTrue(deleted.json()["ok"])
        self.assertEqual(self.client.get("/api/auth/me", headers=headers).status_code, 401)


class RuntimeConfigTests(unittest.TestCase):
    def test_production_rejects_default_or_short_token_secret(self) -> None:
        with self.assertRaises(RuntimeError):
            app_module.validate_runtime_config(
                "production", app_module.DEFAULT_TOKEN_SECRET
            )
        with self.assertRaises(RuntimeError):
            app_module.validate_runtime_config("production", "too-short")

    def test_production_accepts_unique_long_token_secret(self) -> None:
        app_module.validate_runtime_config(
            "production", "unique-production-token-secret-0123456789"
        )


class DocumentExtractionTests(unittest.TestCase):
    def setUp(self) -> None:
        self.client = TestClient(app_module.app)

    def tearDown(self) -> None:
        self.client.close()

    def test_extracts_utf8_txt_for_jd(self) -> None:
        content = "AI 应用开发工程师\n负责 RAG、Agent 工作流与评测体系建设。" * 2
        response = self.client.post(
            "/api/documents/extract?kind=jd",
            files={"file": ("岗位.txt", content.encode("utf-8"), "text/plain")},
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["method"], "文本读取")
        self.assertIn("RAG", payload["text"])
        self.assertFalse(payload["needsReview"])

    def test_extracts_docx_paragraphs_and_tables(self) -> None:
        from docx import Document

        document = Document()
        document.add_paragraph("候选人简历：AI 产品实习生")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "项目"
        table.cell(0, 1).text = "完成 RAG 评测与错误分析"
        stream = io.BytesIO()
        document.save(stream)

        response = self.client.post(
            "/api/documents/extract?kind=resume",
            files={
                "file": (
                    "resume.docx",
                    stream.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["method"], "Word 文本提取")
        self.assertIn("完成 RAG 评测与错误分析", payload["text"])

    def test_extracts_native_pdf_text(self) -> None:
        import fitz

        document = fitz.open()
        page = document.new_page()
        page.insert_text((72, 72), "AI application engineer - Python FastAPI RAG evaluation")
        content = document.tobytes()
        document.close()

        response = self.client.post(
            "/api/documents/extract?kind=jd",
            files={"file": ("job.pdf", content, "application/pdf")},
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["method"], "PDF 文本提取")
        self.assertIn("FastAPI", payload["text"])

    def test_image_ocr_is_marked_for_review(self) -> None:
        with patch.object(
            app_module,
            "_ocr_document_image",
            return_value="项目经历：使用 Python 完成了检索增强生成系统，并记录评测结果。",
        ):
            response = self.client.post(
                "/api/documents/extract?kind=resume",
                files={"file": ("resume.png", b"image-bytes", "image/png")},
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["method"], "图片 OCR")
        self.assertTrue(payload["needsReview"])

    def test_rejects_unsupported_and_oversized_files(self) -> None:
        unsupported = self.client.post(
            "/api/documents/extract?kind=resume",
            files={"file": ("resume.doc", b"legacy word", "application/msword")},
        )
        oversized = self.client.post(
            "/api/documents/extract?kind=resume",
            files={"file": ("resume.txt", b"a" * (app_module.DOCUMENT_MAX_BYTES + 1), "text/plain")},
        )

        self.assertEqual(unsupported.status_code, 422)
        self.assertIn("仅支持", unsupported.json()["detail"])
        self.assertEqual(oversized.status_code, 413)


class DailyHotspotTests(unittest.TestCase):
    def setUp(self) -> None:
        app_module._hotspot_cache.update(
            expires_at=0.0,
            fetched_at=None,
            items=[],
        )

    def test_successful_response_is_normalized(self) -> None:
        response = FakeHotspotResponse(hotspot_payload())
        with patch.object(app_module.httpx, "get", return_value=response):
            result = app_module.get_daily_hotspots(now=100.0)

        self.assertEqual(result["freshness"], "live")
        self.assertEqual(result["items"][0]["sourceName"], "Example News")
        self.assertEqual(result["items"][0]["score"], 79)
        self.assertEqual(result["items"][0]["originalUrl"], "https://example.com/story")
        self.assertEqual(result["items"][0]["providerName"], "AI HOT")
        self.assertEqual(len(result["sources"]), 2)

    def test_ai_insight_feed_is_normalized_with_optional_image(self) -> None:
        responses = [
            FakeHotspotResponse(hotspot_payload()),
            FakeHotspotResponse(insight_payload()),
        ]
        with patch.object(app_module.httpx, "get", side_effect=responses):
            result = app_module.get_daily_hotspots(now=100.0)

        insight_item = next(item for item in result["items"] if item["providerName"] == "AI Insight")
        self.assertEqual(insight_item["category"], "paper")
        self.assertEqual(insight_item["readingUrl"], "https://www.ai-insight.org/news/14469")
        self.assertEqual(insight_item["imageUrl"], "https://images.example.com/paper.png")

    def test_ai_insight_still_loads_when_ai_hot_times_out(self) -> None:
        responses = [
            app_module.httpx.TimeoutException("timeout"),
            FakeHotspotResponse(insight_payload()),
        ]
        with patch.object(app_module.httpx, "get", side_effect=responses):
            result = app_module.get_daily_hotspots(now=100.0)

        self.assertEqual(result["freshness"], "live")
        self.assertEqual(result["items"][0]["providerName"], "AI Insight")

    def test_invalid_original_url_is_not_published(self) -> None:
        response = FakeHotspotResponse(hotspot_payload("javascript:alert(1)"))
        with patch.object(app_module.httpx, "get", return_value=response):
            result = app_module.get_daily_hotspots(now=100.0)

        self.assertEqual(result["freshness"], "unavailable")
        self.assertEqual(result["items"], [])

    def test_provider_timeout_uses_stale_cache(self) -> None:
        cached_item = app_module._normalize_hotspot(hotspot_payload()["items"][0])
        app_module._hotspot_cache.update(
            expires_at=50.0,
            fetched_at="2026-08-05T08:25:11+00:00",
            items=[cached_item],
        )
        with patch.object(
            app_module.httpx,
            "get",
            side_effect=app_module.httpx.TimeoutException("timeout"),
        ):
            result = app_module.get_daily_hotspots(now=100.0)

        self.assertEqual(result["freshness"], "cached")
        self.assertEqual(len(result["items"]), 1)

    def test_fresh_cache_skips_provider_request(self) -> None:
        response = FakeHotspotResponse(hotspot_payload())
        with patch.object(app_module.httpx, "get", return_value=response) as request:
            first = app_module.get_daily_hotspots(now=100.0)
            second = app_module.get_daily_hotspots(now=101.0)

        self.assertEqual(first["items"], second["items"])
        self.assertEqual(request.call_count, 2)


if __name__ == "__main__":
    unittest.main()
