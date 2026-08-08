from __future__ import annotations

import base64
import asyncio
import hashlib
import hmac
import io
import json
import os
import re
import secrets
import sqlite3
import time
import threading
from contextlib import contextmanager
from collections.abc import Iterator
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib.parse import urlsplit

import httpx
from fastapi import BackgroundTasks, Depends, FastAPI, File, Header, HTTPException, Query, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import AliasChoices, BaseModel, Field, model_validator

from .catalog import DAILY_BRIEFS, PROJECTS
from evals.gates import normalize_generator_output, validate_generator_output
from evals.model_client import parse_json_object
from evals.ontology import ontology_prompt_block

ROOT = Path(__file__).resolve().parents[1]


def load_env() -> None:
    for path in (ROOT / ".env", ROOT / "backend" / ".env"):
        if not path.exists():
            continue
        for raw in path.read_text(encoding="utf-8").splitlines():
            line = raw.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


load_env()
DB_PATH = Path(os.getenv("OFFERMAPPING_DB", str(ROOT / "backend" / "offermapping.db")))
APP_ENV = os.getenv("APP_ENV", "development").strip().lower()
DEFAULT_TOKEN_SECRET = "offermapping-local-development-secret"
TOKEN_SECRET = os.getenv("TOKEN_SECRET", DEFAULT_TOKEN_SECRET)
ALLOWED_ORIGINS = [
    origin.strip()
    for origin in os.getenv(
        "ALLOWED_ORIGINS",
        "http://127.0.0.1:4173,http://localhost:4173",
    ).split(",")
    if origin.strip()
]

AI_HOT_API_URL = "https://aihot.virxact.com/api/v1/items"
AI_HOT_SOURCE_URL = "https://aihot.virxact.com/"
AI_INSIGHT_FEED_URL = "https://www.ai-insight.org/feed.json"
AI_INSIGHT_SOURCE_URL = "https://www.ai-insight.org/"
HOTSPOT_CACHE_TTL_SECONDS = 300
HOTSPOT_SOURCE_PAGE_SIZE = 100
DOCUMENT_MAX_BYTES = 10 * 1024 * 1024
DOCUMENT_TEXT_LIMITS = {"jd": 8000, "resume": 10000}
DOCUMENT_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
_hotspot_cache: dict[str, Any] = {
    "expires_at": 0.0,
    "fetched_at": None,
    "window": None,
    "items": [],
}

_analysis_jobs: dict[str, dict[str, Any]] = {}
_analysis_jobs_lock = threading.Lock()
ANALYSIS_JOB_TTL_SECONDS = 30 * 60


class InMemoryRateLimiter:
    """Small single-instance sliding-window limiter for abuse/cost protection."""

    def __init__(self) -> None:
        self._lock = threading.Lock()
        self._hits: dict[str, list[float]] = {}

    def allow(self, key: str, limit: int, window_seconds: int) -> tuple[bool, int]:
        now = time.monotonic()
        with self._lock:
            hits = [stamp for stamp in self._hits.get(key, []) if now - stamp < window_seconds]
            if len(hits) >= limit:
                retry_after = max(1, int(window_seconds - (now - hits[0])))
                self._hits[key] = hits
                return False, retry_after
            hits.append(now)
            self._hits[key] = hits
            return True, 0

    def reset(self) -> None:
        with self._lock:
            self._hits.clear()


rate_limiter = InMemoryRateLimiter()
RATE_LIMIT_WINDOW_SECONDS = max(1, int(os.getenv("RATE_LIMIT_WINDOW_SECONDS", "60")))
RATE_LIMITS = {
    "login": max(1, int(os.getenv("LOGIN_RATE_LIMIT", "10"))),
    "register": max(1, int(os.getenv("REGISTER_RATE_LIMIT", "5"))),
    "analysis": max(1, int(os.getenv("ANALYSIS_RATE_LIMIT", "10"))),
    "documents": max(1, int(os.getenv("DOCUMENT_RATE_LIMIT", "20"))),
    "events": max(1, int(os.getenv("EVENT_RATE_LIMIT", "60"))),
    "feedback": max(1, int(os.getenv("FEEDBACK_RATE_LIMIT", "20"))),
}
COST_WINDOW_SECONDS = max(60, int(os.getenv("COST_WINDOW_SECONDS", str(24 * 60 * 60))))
# Approximate input budget. This protects the demo from accidental model spend;
# production deployments should replace it with a shared quota service.
ANALYSIS_COST_CHAR_LIMIT = max(1000, int(os.getenv("ANALYSIS_COST_CHAR_LIMIT", "300000")))
MODEL_MAX_OUTPUT_TOKENS = max(256, int(os.getenv("MODEL_MAX_OUTPUT_TOKENS", "1800")))


class WeightedBudget:
    def __init__(self) -> None:
        self._lock = threading.Lock()
        self._usage: dict[str, list[tuple[float, int]]] = {}

    def consume(self, key: str, amount: int, limit: int, window_seconds: int) -> tuple[bool, int]:
        now = time.monotonic()
        with self._lock:
            entries = [entry for entry in self._usage.get(key, []) if now - entry[0] < window_seconds]
            used = sum(value for _, value in entries)
            if used + amount > limit:
                retry_after = max(1, int(window_seconds - (now - entries[0][0]))) if entries else window_seconds
                self._usage[key] = entries
                return False, retry_after
            entries.append((now, amount))
            self._usage[key] = entries
            return True, 0

    def reset(self) -> None:
        with self._lock:
            self._usage.clear()


analysis_budget = WeightedBudget()


def _safe_http_url(value: Any) -> str | None:
    if not isinstance(value, str):
        return None
    candidate = value.strip()
    parsed = urlsplit(candidate)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        return None
    return candidate


def _clean_hotspot_text(value: Any, limit: int) -> str:
    if not isinstance(value, str):
        return ""
    text = re.sub(r"\s+", " ", value).strip()
    return text if len(text) <= limit else f"{text[: limit - 1].rstrip()}…"


_document_ocr_engine: Any = None


def _normalize_document_text(value: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in value.replace("\x00", "").splitlines()]
    normalized: list[str] = []
    previous_blank = False
    for line in lines:
        is_blank = not line
        if is_blank and previous_blank:
            continue
        normalized.append(line)
        previous_blank = is_blank
    return "\n".join(normalized).strip()


def _ocr_document_image(content: bytes) -> str:
    global _document_ocr_engine
    try:
        from rapidocr import EngineType, RapidOCR
    except ImportError as error:
        raise RuntimeError("图片文字识别组件尚未安装。") from error
    if _document_ocr_engine is None:
        _document_ocr_engine = RapidOCR(
            params={
                "Det.engine_type": EngineType.OPENVINO,
                "Cls.engine_type": EngineType.OPENVINO,
                "Rec.engine_type": EngineType.OPENVINO,
            }
        )
    result = _document_ocr_engine(content)
    texts = getattr(result, "txts", None) or ()
    if not texts:
        return ""
    return "\n".join(text.strip() for text in texts if isinstance(text, str) and text.strip())


def _extract_pdf_document(content: bytes) -> tuple[str, bool]:
    try:
        import fitz
    except ImportError as error:
        raise RuntimeError("PDF 解析组件尚未安装。") from error
    text_pages: list[str] = []
    used_ocr = False
    with fitz.open(stream=content, filetype="pdf") as document:
        if document.page_count > 40:
            raise ValueError("PDF 最多支持 40 页，请拆分后重新上传。")
        for page in document:
            page_text = page.get_text("text").strip()
            if len(page_text) < 20:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
                page_text = _ocr_document_image(pixmap.tobytes("png"))
                used_ocr = True
            if page_text:
                text_pages.append(page_text)
    return "\n\n".join(text_pages), used_ocr


def _extract_docx_document(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError("Word 解析组件尚未安装。") from error
    document = Document(io.BytesIO(content))
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def _extract_uploaded_document(filename: str, content: bytes) -> tuple[str, str, bool]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".txt":
        for encoding in ("utf-8-sig", "gb18030", "utf-16"):
            try:
                return content.decode(encoding), "文本读取", False
            except UnicodeDecodeError:
                continue
        raise ValueError("无法识别 TXT 文件编码，请另存为 UTF-8 后重试。")
    if suffix == ".pdf":
        text, used_ocr = _extract_pdf_document(content)
        return text, "PDF + OCR" if used_ocr else "PDF 文本提取", used_ocr
    if suffix == ".docx":
        return _extract_docx_document(content), "Word 文本提取", False
    if suffix in DOCUMENT_IMAGE_SUFFIXES:
        return _ocr_document_image(content), "图片 OCR", True
    raise ValueError("仅支持 TXT、PDF、DOCX、PNG、JPG、WEBP、BMP 和 TIFF 文件。")


def _normalize_hotspot(item: Any) -> dict[str, Any] | None:
    if not isinstance(item, dict):
        return None
    title = _clean_hotspot_text(item.get("title"), 180)
    links = item.get("links") if isinstance(item.get("links"), dict) else {}
    original_url = _safe_http_url(links.get("original"))
    aihot_url = _safe_http_url(links.get("aihot"))
    if not title or not original_url or not aihot_url:
        return None

    source = item.get("source") if isinstance(item.get("source"), dict) else {}
    try:
        score = max(0, min(100, int(item.get("score", 0))))
    except (TypeError, ValueError):
        score = 0

    return {
        "id": _clean_hotspot_text(item.get("id"), 120) or aihot_url,
        "title": title,
        "summary": _clean_hotspot_text(item.get("summary"), 360),
        "category": _clean_hotspot_text(item.get("category"), 60) or "ai-industry",
        "score": score,
        "sourceName": _clean_hotspot_text(source.get("name"), 80) or "原始来源",
        "originalUrl": original_url,
        "aihotUrl": aihot_url,
        "readingUrl": aihot_url,
        "providerName": "AI HOT",
        "imageUrl": None,
        "publishedAt": _clean_hotspot_text(item.get("publishedAt"), 64),
    }


def _normalize_insight_category(item: dict[str, Any]) -> str:
    tags = item.get("tags") if isinstance(item.get("tags"), list) else []
    tag = _clean_hotspot_text(tags[0], 40) if tags else ""
    return {
        "研究": "paper",
        "大模型": "ai-models",
        "产品发布": "ai-products",
        "产品观察": "ai-products",
        "芯片": "ai-infrastructure",
        "机器人": "ai-products",
        "安全": "industry",
    }.get(tag, "ai-industry")


def _normalize_insight_item(item: Any) -> dict[str, Any] | None:
    if not isinstance(item, dict):
        return None
    title = _clean_hotspot_text(item.get("title"), 180)
    original_url = _safe_http_url(item.get("external_url") or item.get("url"))
    reading_url = _safe_http_url(item.get("id"))
    if not title or not original_url or not reading_url:
        return None

    authors = item.get("authors") if isinstance(item.get("authors"), list) else []
    author = authors[0] if authors and isinstance(authors[0], dict) else {}
    summary = item.get("summary") or item.get("content_text")

    return {
        "id": f"ai-insight:{reading_url}",
        "title": title,
        "summary": _clean_hotspot_text(summary, 360),
        "category": _normalize_insight_category(item),
        "score": 0,
        "sourceName": _clean_hotspot_text(author.get("name"), 80) or "原始来源",
        "originalUrl": original_url,
        "aihotUrl": reading_url,
        "readingUrl": reading_url,
        "providerName": "AI Insight",
        "imageUrl": _safe_http_url(item.get("image")),
        "publishedAt": _clean_hotspot_text(item.get("date_published"), 64),
    }


def _fetch_ai_hot_items(window: str = "24h") -> list[dict[str, Any]]:
    raw_items: list[Any] = []
    cursor: str | None = None
    seen_cursors: set[str] = set()
    while True:
        params: dict[str, Any] = {"mode": "selected", "window": window, "limit": HOTSPOT_SOURCE_PAGE_SIZE}
        if cursor:
            params["cursor"] = cursor
        response = httpx.get(
            AI_HOT_API_URL,
            params=params,
            headers={"User-Agent": "OfferMapping-AI/0.2"},
            timeout=8.0,
        )
        response.raise_for_status()
        payload = response.json()
        if not isinstance(payload, dict):
            break
        raw_items.extend(payload.get("items", []) if isinstance(payload.get("items"), list) else [])
        page = payload.get("page") if isinstance(payload.get("page"), dict) else {}
        next_cursor = page.get("nextCursor") if isinstance(page.get("nextCursor"), str) else None
        if not page.get("hasMore") or not next_cursor or next_cursor in seen_cursors:
            break
        seen_cursors.add(next_cursor)
        cursor = next_cursor
    return [normalized for raw in raw_items if (normalized := _normalize_hotspot(raw))]


def _fetch_ai_insight_items() -> list[dict[str, Any]]:
    response = httpx.get(
        AI_INSIGHT_FEED_URL,
        headers={"User-Agent": "OfferMapping-AI/0.2"},
        timeout=8.0,
    )
    response.raise_for_status()
    payload = response.json()
    raw_items = payload.get("items", []) if isinstance(payload, dict) else []
    return [normalized for raw in raw_items if (normalized := _normalize_insight_item(raw))]


HOTSPOT_JOB_TERMS = (
    "agent",
    "智能体",
    "model",
    "模型",
    "eval",
    "评测",
    "benchmark",
    "基准",
    "coding",
    "编程",
    "security",
    "安全",
    "deployment",
    "部署",
    "rag",
    "检索",
    "api",
    "开源",
    "工具",
    "workflow",
    "工作流",
)


def _hotspot_relevance_score(item: dict[str, Any]) -> int:
    text = f"{item.get('title', '')} {item.get('summary', '')}".lower()
    return sum(1 for term in HOTSPOT_JOB_TERMS if term in text)


def _merge_hotspot_items(*groups: list[dict[str, Any]]) -> list[dict[str, Any]]:
    unique: dict[str, dict[str, Any]] = {}
    for item in (item for group in groups for item in group):
        unique.setdefault(item["originalUrl"], item)
    return sorted(
        unique.values(),
        key=lambda item: (
            item.get("publishedAt", "")[:10],
            _hotspot_relevance_score(item),
            item.get("score", 0),
            item.get("publishedAt", ""),
        ),
        reverse=True,
    )


def _select_weekly_hotspot_items(*groups: list[dict[str, Any]]) -> list[dict[str, Any]]:
    unique: dict[str, dict[str, Any]] = {}
    for item in (item for group in groups for item in group):
        unique.setdefault(item["originalUrl"], item)

    by_day: dict[str, list[dict[str, Any]]] = {}
    for item in sorted(
        unique.values(),
        key=lambda candidate: (
            candidate.get("publishedAt", "")[:10],
            _hotspot_relevance_score(candidate),
            candidate.get("score", 0),
            candidate.get("publishedAt", ""),
        ),
        reverse=True,
    ):
        day = item.get("publishedAt", "")[:10] or "unknown"
        by_day.setdefault(day, []).append(item)

    days = sorted(by_day, reverse=True)
    selected: list[dict[str, Any]] = []
    while any(by_day[day] for day in days):
        for day in days:
            if by_day[day]:
                selected.append(by_day[day].pop(0))
    return selected


def get_daily_hotspots(now: float | None = None, window: str = "24h") -> dict[str, Any]:
    current_time = time.monotonic() if now is None else now
    window_key = "7d" if window == "7d" else "24h"
    cached_items = _hotspot_cache["items"] if _hotspot_cache.get("window") in {None, window_key} else []
    daily_url = f"{AI_HOT_SOURCE_URL}daily/{datetime.now().strftime('%Y-%m-%d')}"
    sources = [
        {"name": "AI HOT", "url": AI_HOT_SOURCE_URL},
        {"name": "AI Insight", "url": AI_INSIGHT_SOURCE_URL},
    ]
    if cached_items and current_time < _hotspot_cache["expires_at"] and _hotspot_cache.get("window") == window_key:
        return {
            "windowKey": window_key,
            "window": "过去 7 天" if window_key == "7d" else "过去 24 小时",
            "freshness": "live",
            "fetchedAt": _hotspot_cache["fetched_at"],
            "dailyUrl": daily_url,
            "source": {"name": "AI HOT", "url": AI_HOT_SOURCE_URL},
            "sources": sources,
            "items": cached_items,
        }

    provider_items: list[list[dict[str, Any]]] = []
    for fetch_provider in (lambda: _fetch_ai_hot_items(window_key), _fetch_ai_insight_items):
        try:
            fetched_items = fetch_provider()
            if fetched_items:
                provider_items.append(fetched_items)
        except (httpx.HTTPError, ValueError, TypeError):
            continue

    if window_key == "7d" and len(provider_items) > 1:
        # Interleave by publication date so the weekly pool keeps its full
        # history instead of being dominated by the freshest feed.
        items = _select_weekly_hotspot_items(*provider_items)
    else:
        items = _merge_hotspot_items(*provider_items)
    if items:
        fetched_at = datetime.now(timezone.utc).isoformat()
        _hotspot_cache.update(
            expires_at=current_time + HOTSPOT_CACHE_TTL_SECONDS,
            fetched_at=fetched_at,
            window=window_key,
            items=items,
        )
        freshness = "live"
    else:
        items = cached_items
        fetched_at = _hotspot_cache["fetched_at"]
        freshness = "cached" if items else "unavailable"

    return {
        "windowKey": window_key,
        "window": "过去 7 天" if window_key == "7d" else "过去 24 小时",
        "freshness": freshness,
        "fetchedAt": fetched_at,
        "dailyUrl": daily_url,
        "source": {"name": "AI HOT", "url": AI_HOT_SOURCE_URL},
        "sources": sources,
        "items": items,
    }


def validate_runtime_config(app_env: str, token_secret: str) -> None:
    if app_env == "production" and (
        token_secret == DEFAULT_TOKEN_SECRET or len(token_secret) < 32
    ):
        raise RuntimeError(
            "Production requires TOKEN_SECRET to be set to a unique value of at least 32 characters."
        )


validate_runtime_config(APP_ENV, TOKEN_SECRET)

app = FastAPI(title="OfferMapping AI", version="0.2.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def _request_id(request: Request) -> str:
    value = request.headers.get("X-Request-ID", "").strip()
    # Accept caller supplied IDs for trace correlation, but cap the value and
    # reject control characters so it is safe to put in logs and headers.
    if value and len(value) <= 128 and all(ord(char) >= 32 for char in value):
        return value
    return secrets.token_hex(16)


def _error_fingerprint(path: str, status_code: int, detail: Any) -> str:
    normalized = re.sub(r"\d+", "#", str(detail).lower())[:240]
    return hashlib.sha256(f"{path}|{status_code}|{normalized}".encode()).hexdigest()[:24]


def record_error(request: Request, status_code: int, detail: Any, request_id: str) -> None:
    """Aggregate failures without storing request bodies or sensitive data."""
    try:
        fingerprint = _error_fingerprint(request.url.path, status_code, detail)
        with db() as connection:
            connection.execute(
                """INSERT INTO error_events(fingerprint,path,status_code,count,last_message,
                   first_seen,last_seen,last_request_id) VALUES(?,?,?,?,?,?,?,?)
                   ON CONFLICT(fingerprint) DO UPDATE SET count=count+1,
                   last_message=excluded.last_message,last_seen=excluded.last_seen,
                   last_request_id=excluded.last_request_id""",
                (
                    fingerprint,
                    request.url.path,
                    status_code,
                    1,
                    str(detail)[:240],
                    now_iso(),
                    now_iso(),
                    request_id,
                ),
            )
    except Exception:
        # Error reporting must never mask the original response.
        return


@app.middleware("http")
async def request_context(request: Request, call_next):
    request_id = _request_id(request)
    request.state.request_id = request_id
    try:
        response = await call_next(request)
    except Exception as error:
        record_error(request, 500, type(error).__name__, request_id)
        request.state.error_recorded = True
        response = JSONResponse(
            status_code=500,
            content={"detail": "服务暂时不可用，请稍后重试。", "requestId": request_id},
        )
    if response.status_code >= 400 and not getattr(request.state, "error_recorded", False):
        record_error(request, response.status_code, f"http_{response.status_code}", request_id)
    response.headers["X-Request-ID"] = request_id
    return response


@app.exception_handler(HTTPException)
async def handle_http_error(request: Request, exc: HTTPException) -> JSONResponse:
    request_id = getattr(request.state, "request_id", _request_id(request))
    record_error(request, exc.status_code, exc.detail, request_id)
    request.state.error_recorded = True
    headers = dict(exc.headers or {})
    headers["X-Request-ID"] = request_id
    return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail, "requestId": request_id}, headers=headers)


def enforce_rate_limit(request: Request, scope: str, user: dict[str, Any] | None = None) -> None:
    identity = str(user.get("sub")) if user else (request.client.host if request.client else "unknown")
    allowed, retry_after = rate_limiter.allow(f"{scope}:{identity}", RATE_LIMITS[scope], RATE_LIMIT_WINDOW_SECONDS)
    if not allowed:
        raise HTTPException(
            status_code=429,
            detail="请求过于频繁，请稍后再试。",
            headers={"Retry-After": str(retry_after)},
        )


def enforce_analysis_budget(request: Request, payload: "AnalysisPayload", user: dict[str, Any] | None) -> None:
    identity = str(user.get("sub")) if user else (request.client.host if request.client else "unknown")
    chars = len(payload.jd) + len(payload.resume)
    allowed, retry_after = analysis_budget.consume(
        f"analysis-cost:{identity}", chars, ANALYSIS_COST_CHAR_LIMIT, COST_WINDOW_SECONDS
    )
    if not allowed:
        raise HTTPException(
            status_code=429,
            detail="已达到本周期分析额度，请稍后再试。",
            headers={"Retry-After": str(retry_after)},
        )


@contextmanager
def db() -> Iterator[sqlite3.Connection]:
    connection = sqlite3.connect(DB_PATH)
    connection.row_factory = sqlite3.Row
    try:
        with connection:
            yield connection
    finally:
        connection.close()


def init_db() -> None:
    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    with db() as connection:
        connection.executescript(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                email TEXT NOT NULL UNIQUE,
                password_hash TEXT NOT NULL,
                created_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS analyses (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                role TEXT NOT NULL,
                score INTEGER NOT NULL,
                jd_raw TEXT NOT NULL,
                resume_raw TEXT NOT NULL,
                result_json TEXT NOT NULL,
                source TEXT NOT NULL,
                created_at TEXT NOT NULL,
                FOREIGN KEY(user_id) REFERENCES users(id)
            );
            CREATE TABLE IF NOT EXISTS model_runs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                analysis_id INTEGER,
                role TEXT NOT NULL,
                model TEXT NOT NULL,
                latency_ms INTEGER NOT NULL,
                input_tokens INTEGER NOT NULL DEFAULT 0,
                output_tokens INTEGER NOT NULL DEFAULT 0,
                schema_ok INTEGER NOT NULL DEFAULT 0,
                error TEXT,
                created_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS events (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                analysis_id INTEGER,
                event TEXT NOT NULL,
                metadata_json TEXT NOT NULL DEFAULT '{}',
                created_at TEXT NOT NULL,
                FOREIGN KEY(user_id) REFERENCES users(id),
                FOREIGN KEY(analysis_id) REFERENCES analyses(id)
            );
            CREATE TABLE IF NOT EXISTS feedback (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                analysis_id INTEGER,
                positive INTEGER,
                comment TEXT,
                created_at TEXT NOT NULL,
                FOREIGN KEY(user_id) REFERENCES users(id),
                FOREIGN KEY(analysis_id) REFERENCES analyses(id)
            );
            CREATE TABLE IF NOT EXISTS error_events (
                fingerprint TEXT PRIMARY KEY,
                path TEXT NOT NULL,
                status_code INTEGER NOT NULL,
                count INTEGER NOT NULL DEFAULT 0,
                last_message TEXT NOT NULL,
                first_seen TEXT NOT NULL,
                last_seen TEXT NOT NULL,
                last_request_id TEXT NOT NULL
            );
            CREATE INDEX IF NOT EXISTS idx_events_analysis_id ON events(analysis_id);
            CREATE INDEX IF NOT EXISTS idx_feedback_analysis_id ON feedback(analysis_id);
            """
        )


init_db()


class AuthPayload(BaseModel):
    email: str
    password: str = Field(min_length=8, max_length=128)


class AnalysisPayload(BaseModel):
    jd: str = Field(min_length=50, max_length=8000)
    resume: str = Field(min_length=100, max_length=10000)


class EventPayload(BaseModel):
    analysis_id: int | None = Field(default=None, ge=1, validation_alias=AliasChoices("analysis_id", "analysisId"))
    event: str = Field(min_length=2, max_length=64, pattern=r"^[a-z][a-z0-9_.-]+$")
    metadata: dict[str, Any] = Field(default_factory=dict)


class FeedbackPayload(BaseModel):
    analysis_id: int | None = Field(default=None, ge=1, validation_alias=AliasChoices("analysis_id", "analysisId"))
    positive: bool | None = None
    rating: str | None = Field(default=None, pattern="^(up|down)$")
    comment: str | None = Field(default=None, max_length=2000)

    @model_validator(mode="after")
    def normalize_rating(self) -> "FeedbackPayload":
        if self.positive is None and self.rating is not None:
            self.positive = self.rating == "up"
        if self.positive is None and not self.comment:
            raise ValueError("feedback requires positive/rating or comment")
        return self


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def hash_password(password: str) -> str:
    salt = secrets.token_bytes(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt, 240_000)
    return f"pbkdf2_sha256$240000${base64.urlsafe_b64encode(salt).decode()}${base64.urlsafe_b64encode(digest).decode()}"


def verify_password(password: str, encoded: str) -> bool:
    try:
        _, iterations, salt_text, digest_text = encoded.split("$", 3)
        salt = base64.urlsafe_b64decode(salt_text)
        expected = base64.urlsafe_b64decode(digest_text)
        actual = hashlib.pbkdf2_hmac("sha256", password.encode(), salt, int(iterations))
        return hmac.compare_digest(actual, expected)
    except (ValueError, TypeError):
        return False


def b64url(data: bytes) -> str:
    return base64.urlsafe_b64encode(data).decode().rstrip("=")


def make_token(user_id: int, email: str) -> str:
    payload = {"sub": user_id, "email": email, "exp": int(time.time()) + 7 * 86400}
    body = b64url(json.dumps(payload, separators=(",", ":")).encode())
    signature = b64url(hmac.new(TOKEN_SECRET.encode(), body.encode(), hashlib.sha256).digest())
    return f"{body}.{signature}"


def decode_token(token: str) -> dict[str, Any] | None:
    try:
        body, signature = token.split(".", 1)
        expected = b64url(hmac.new(TOKEN_SECRET.encode(), body.encode(), hashlib.sha256).digest())
        if not hmac.compare_digest(signature, expected):
            return None
        padded = body + "=" * (-len(body) % 4)
        payload = json.loads(base64.urlsafe_b64decode(padded))
        return payload if payload.get("exp", 0) > time.time() else None
    except (ValueError, json.JSONDecodeError):
        return None


def optional_user(authorization: str | None) -> dict[str, Any] | None:
    if not authorization or not authorization.lower().startswith("bearer "):
        return None
    return decode_token(authorization.split(" ", 1)[1])


def require_user(authorization: str | None = Header(default=None)) -> dict[str, Any]:
    user = optional_user(authorization)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录后查看历史记录。")
    with db() as connection:
        exists = connection.execute("SELECT 1 FROM users WHERE id=?", (user["sub"],)).fetchone()
    if not exists:
        raise HTTPException(status_code=401, detail="登录状态已失效，请重新登录。")
    return user


def model_config(role: str) -> dict[str, str]:
    prefix = role.upper()
    profile_name = os.getenv(f"{prefix}_PROFILE", "")
    registry_path = Path(os.getenv("MODEL_REGISTRY_PATH", ROOT / "backend" / "models.json"))
    if profile_name and registry_path.exists():
        try:
            profile = json.loads(registry_path.read_text(encoding="utf-8")).get(profile_name, {})
            api_key_env = profile.get("api_key_env", "")
            return {
                "base_url": str(profile.get("base_url", "")).rstrip("/"),
                "api_key": os.getenv(api_key_env, "") if api_key_env else "",
                "model": str(profile.get("model", "")),
                "profile": profile_name,
            }
        except (OSError, json.JSONDecodeError):
            pass
    return {
        "base_url": os.getenv(f"{prefix}_BASE_URL", "").rstrip("/"),
        "api_key": os.getenv(f"{prefix}_API_KEY", ""),
        "model": os.getenv(f"{prefix}_MODEL", ""),
        "profile": "",
    }


def json_from_text(text: str) -> dict[str, Any]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned, flags=re.I | re.S)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        start, end = cleaned.find("{"), cleaned.rfind("}")
        if start >= 0 and end > start:
            return json.loads(cleaned[start : end + 1])
        raise


async def call_model(role: str, system: str, user: str, temperature: float) -> tuple[dict[str, Any] | None, dict[str, Any]]:
    config = model_config(role)
    meta = {"role": role, "model": config["model"] or "not-configured", "latency_ms": 0, "input_tokens": 0, "output_tokens": 0, "schema_ok": False, "error": None}
    if not all(config[key] for key in ("base_url", "api_key", "model")):
        meta["error"] = "model_not_configured"
        return None, meta
    payload = {
        "model": config["model"],
        "messages": [{"role": "system", "content": system}, {"role": "user", "content": user}],
        "temperature": temperature,
        "max_tokens": MODEL_MAX_OUTPUT_TOKENS,
        "response_format": {"type": "json_object"},
    }
    started = time.perf_counter()
    request_headers = {
        "Authorization": f"Bearer {config['api_key']}",
        "Content-Type": "application/json",
        "User-Agent": os.getenv("MODEL_USER_AGENT", "OfferMapping-Backend/0.2"),
    }
    last_content = ""
    try:
        async with httpx.AsyncClient(timeout=90) as client:
            for attempt in range(2):
                request_payload = dict(payload)
                response = await client.post(
                    f"{config['base_url']}/chat/completions",
                    headers=request_headers,
                    json=request_payload,
                )
                if response.status_code == 400 and "response_format" in request_payload:
                    request_payload.pop("response_format", None)
                    response = await client.post(
                        f"{config['base_url']}/chat/completions",
                        headers=request_headers,
                        json=request_payload,
                    )
                response.raise_for_status()
                raw = response.json()
                content = raw["choices"][0]["message"]["content"]
                last_content = content if isinstance(content, str) else str(content)
                result = parse_json_object(last_content)
                usage = raw.get("usage") or {}
                meta.update(
                    latency_ms=round((time.perf_counter() - started) * 1000),
                    input_tokens=usage.get("prompt_tokens", 0),
                    output_tokens=usage.get("completion_tokens", 0),
                    schema_ok=True,
                )
                return result, meta
    except Exception as exc:  # external providers vary; the UI receives a safe fallback
        error_text = f"{type(exc).__name__}: {str(exc)[:220]}"
        if last_content and isinstance(exc, json.JSONDecodeError):
            error_text = f"{error_text} raw={last_content[:120]}"
        meta.update(latency_ms=round((time.perf_counter() - started) * 1000), error=error_text)
        return None, meta


SKILL_DEFS = [
    ("python", "Python", [r"python"]),
    ("rag", "RAG", [r"\brag\b", r"检索增强"]),
    ("llm", "大模型应用", [r"\bllm\b", r"大模型", r"语言模型"]),
    ("agent", "Agent / Function Calling", [r"agent", r"智能体", r"function calling", r"工具调用"]),
    ("evals", "LLM 评估", [r"evals?", r"评测", r"评估体系", r"质量评估", r"对比实验", r"提示词.{0,8}对比", r"失败案例"]),
    ("fastapi", "FastAPI", [r"fastapi"]),
    ("docker", "Docker", [r"docker", r"容器化"]),
    ("sql", "SQL / 数据库", [r"\bsql\b", r"sqlite", r"postgres", r"mysql", r"数据库"]),
    ("vector", "向量数据库", [r"pgvector", r"milvus", r"faiss", r"向量数据库"]),
    ("prompt", "Prompt Engineering", [r"prompt", r"提示词"]),
    ("deployment", "部署与服务化", [r"部署", r"服务化", r"云服务"]),
]


_NEGATED_EVIDENCE = re.compile(
    r"(?:不要求|无需|无须|不需要|不具备|不熟悉|不了解|没有(?:做过|使用|接触|(?:线上)?部署|项目|经验|实际使用)|没(?:做过|使用|接触)"
    r"|未(?:使用|做过|接触|参与)|尚未|未曾|不参与|不涉及|不负责)",
    re.I,
)
_NEGATED_BEFORE_TARGET = re.compile(r"(?:不要求|无需|无须|不需要|不具备|不熟悉|不了解|没有|没|未|尚未|不参与|不涉及|不负责)(?:线上|实际|任何)?\s*$", re.I)
_UNTRUSTED_TEXT = re.compile(
    r"(?:以下(?:内容|文字)是(?:测试|示例)文本|测试文本|提示注入|prompt\s*injection|"
    r"忽略(?:岗位|上文|以上)(?:要求|指令)|输出我(?:精通|擅长)|请(?:直接)?输出)",
    re.I,
)
_PROJECT_EVIDENCE = re.compile(
    r"(?:项目|课程项目|个人项目|毕业设计|实习中|工作中|使用.+?(?:完成|实现|开发|编写|构建|搭建|制作|部署|调用|设计|维护|查询|清洗|分析|测试)|"
    r"(?:完成|实现|开发|编写|构建|搭建|制作|部署|调用|设计|维护|负责|上线|排查|训练|分析|测试).{0,24})",
    re.I,
)
_LISTED_ONLY = re.compile(
    r"(?:技能(?:栏|包括|清单)|列出|了解|熟悉|学习过|课程(?:作业|考试|项目)?|考试|教程|安装过|只安装|听说过|接触过)",
    re.I,
)
_WEAK_SOURCE = re.compile(r"(?:参加过|分享会|阅读过|读过|观看过|看过|文章|文档)", re.I)
_TUTORIAL_OR_INSTALL = re.compile(r"(?:教程|安装过|只安装|只完成过.+?(?:教程|入门))", re.I)
_COURSE_KNOWLEDGE = re.compile(r"(?:入门课程|课程考试|课程培训|培训课程)", re.I)


def _sentences(text: str) -> list[str]:
    """Split source text without allowing a neighboring sentence to become evidence."""
    return [part.strip() for part in re.split(r"[。！？!?\n]", text) if part.strip()]


def _pattern_is_negated(sentence: str, pattern: str) -> bool:
    """Treat a negation close to the matched term as absence, not as a skill claim."""
    for match in re.finditer(pattern, sentence, re.I):
        prefix = sentence[max(0, match.start() - 24) : match.start()]
        suffix = sentence[match.end() : min(len(sentence), match.end() + 16)]
        if _NEGATED_EVIDENCE.search(prefix) or _NEGATED_EVIDENCE.search(suffix):
            return True
    return False


def _is_usable_match(sentence: str, patterns: list[str]) -> bool:
    if _UNTRUSTED_TEXT.search(sentence):
        return False

    def usable_occurrence(pattern: str) -> bool:
        for match in re.finditer(pattern, sentence, re.I):
            prefix = re.split(r"[，,；;]", sentence[max(0, match.start() - 24) : match.start()])[-1]
            suffix = re.split(r"[，,；;]", sentence[match.end() : min(len(sentence), match.end() + 16)])[0]
            if not (
                _NEGATED_EVIDENCE.search(prefix)
                or _NEGATED_BEFORE_TARGET.search(prefix)
                or _NEGATED_EVIDENCE.search(suffix)
            ):
                return True
        return False

    return any(
        usable_occurrence(pattern)
        for pattern in patterns
    )


def _resume_evidence(quote: str) -> str:
    """Classify only the quoted local sentence, conservatively."""
    if not quote or _UNTRUSTED_TEXT.search(quote):
        return "missing"
    # Explicitly weak sources are not project proof. Reading or an event is not a
    # claim of experience; course exams/tutorials may support listed-only knowledge.
    if _WEAK_SOURCE.search(quote) and not re.search(r"课程(?:项目|作业)|实习|工作", quote):
        return "missing"
    if _TUTORIAL_OR_INSTALL.search(quote) or _COURSE_KNOWLEDGE.search(quote):
        return "listed-only"
    if _PROJECT_EVIDENCE.search(quote) and not _NEGATED_EVIDENCE.search(quote):
        return "project-backed"
    return "listed-only"


def sentence_match(text: str, patterns: list[str], evidence_first: bool = False) -> str:
    matches = [sentence for sentence in _sentences(text) if _is_usable_match(sentence, patterns)]
    if evidence_first:
        evidence = next((item for item in matches if _resume_evidence(item) == "project-backed"), None)
        if evidence:
            return evidence[:100]
    return matches[0][:100] if matches else ""


def infer_job_family(jd: str) -> str:
    if re.search(r"产品经理|产品设计|用户研究", jd):
        return "ai_product"
    if re.search(r"算法|训练|微调|推理优化", jd):
        return "algorithm"
    if re.search(r"数据分析|数据工程|数据科学", jd):
        return "data"
    return "ai_app_dev"


def infer_role(jd: str, family: str) -> str:
    for line in jd.splitlines():
        if re.search(r"工程师|产品经理|算法|数据分析|开发", line):
            return line.strip()[:32]
    return {"ai_product": "AI 产品经理", "algorithm": "大模型算法工程师", "data": "AI 数据工程师", "ai_app_dev": "AI 应用开发工程师"}[family]


def background_assets(resume: str) -> list[str]:
    assets = []
    mapping = [
        (r"金融|财务|会计|证券|银行", "金融与业务分析背景"),
        (r"传媒|内容|新闻|运营|广告", "内容与传播背景"),
        (r"教育|教学|课程", "教育与学习场景经验"),
        (r"计算机|软件|信息工程|开发", "软件开发基础"),
        (r"医疗|医学|生物", "医疗与生命科学背景"),
    ]
    for pattern, label in mapping:
        if re.search(pattern, resume, re.I):
            assets.append(label)
    return assets or ["已有学习与项目经历"]


def local_extract(jd: str, resume: str) -> dict[str, Any]:
    nice_words = re.compile(r"加分|优先|了解|熟悉更佳|nice.?to.?have", re.I)
    skills = []
    for key, name, patterns in SKILL_DEFS:
        jd_quote = sentence_match(jd, patterns)
        if not jd_quote:
            continue
        resume_quote = sentence_match(resume, patterns, True)
        evidence = _resume_evidence(resume_quote)
        if evidence == "missing":
            resume_quote = ""
        skills.append({"key": key, "name": name, "priority": "nice" if nice_words.search(jd_quote) else "must", "jd_quote": jd_quote, "resume_quote": resume_quote, "evidence": evidence})
    if len(skills) < 4:
        for key, name, patterns in SKILL_DEFS[:6]:
            if any(item["key"] == key for item in skills):
                continue
            jd_fallback_quote = sentence_match(jd, patterns)
            if not jd_fallback_quote:
                continue
            resume_quote = sentence_match(resume, patterns, True)
            evidence = _resume_evidence(resume_quote)
            if evidence == "missing":
                resume_quote = ""
            skills.append({"key": key, "name": name, "priority": "must" if len(skills) < 4 else "nice", "jd_quote": jd_fallback_quote, "resume_quote": resume_quote, "evidence": evidence})
    family = infer_job_family(jd)
    return {"job_family": family, "role": infer_role(jd, family), "background_assets": background_assets(resume), "skills": skills}


def normalize_model_extract(raw: dict[str, Any], jd: str, resume: str) -> dict[str, Any]:
    fallback = local_extract(jd, resume)
    normalized = []
    for item in raw.get("skills", []):
        jd_quote = str(item.get("jd_quote", "")).strip()
        resume_quote = str(item.get("resume_quote", "")).strip()
        if jd_quote and jd_quote not in jd:
            continue
        if resume_quote and resume_quote not in resume:
            resume_quote = ""
        evidence = item.get("evidence")
        if evidence not in {"project-backed", "listed-only", "missing"}:
            evidence = "project-backed" if resume_quote else "missing"
        normalized.append({
            "key": re.sub(r"[^a-z0-9_]+", "_", str(item.get("key") or item.get("name") or "skill").lower()).strip("_"),
            "name": str(item.get("name") or "未命名技能")[:40],
            "priority": "nice" if item.get("priority") == "nice" else "must",
            "jd_quote": jd_quote or "岗位原文未提供",
            "resume_quote": resume_quote,
            "evidence": evidence if resume_quote else "missing",
        })
    if len(normalized) < 3:
        return fallback
    family = raw.get("job_family") if raw.get("job_family") in {"algorithm", "ai_app_dev", "ai_product", "data"} else fallback["job_family"]
    return {
        "job_family": family,
        "role": str(raw.get("role") or fallback["role"])[:40],
        "background_assets": [str(value)[:60] for value in raw.get("background_assets", [])[:5]] or fallback["background_assets"],
        "skills": normalized,
    }


def time_for_skill(key: str) -> str:
    if key in {"docker", "fastapi", "prompt", "deployment"}:
        return "2–3 天"
    if key in {"rag", "agent", "sql", "vector"}:
        return "4–7 天"
    return "1–2 周"


def score_skills(skills: list[dict[str, Any]]) -> tuple[int, list[dict[str, int]], str]:
    value = {"project-backed": 1.0, "listed-only": 0.5, "missing": 0.0}
    must = [item for item in skills if item["priority"] == "must"]
    nice = [item for item in skills if item["priority"] == "nice"]
    must_rate = sum(value[item["evidence"]] for item in must) / max(len(must), 1)
    nice_rate = sum(value[item["evidence"]] for item in nice) / max(len(nice), 1)
    project_rate = sum(1 for item in skills if item["evidence"] == "project-backed") / max(len(skills), 1)
    score = round(must_rate * 60 + nice_rate * 20 + project_rate * 20)
    dimensions = [
        {"label": "硬技能", "score": round(must_rate * 40), "max": 40},
        {"label": "项目证据", "score": round(project_rate * 40), "max": 40},
        {"label": "领域匹配", "score": min(20, round((must_rate * 0.6 + nice_rate * 0.4) * 20)), "max": 20},
    ]
    missing = [item["name"] for item in skills if item["evidence"] == "missing"]
    primary = missing[0] if missing else "项目深度"
    return score, dimensions, primary


def select_focus_skill_keys(skills: list[dict[str, Any]], limit: int = 3) -> list[str]:
    """Select the most useful evidence gaps for Generator recommendations."""
    tiers = (
        ("must", {"missing"}),
        ("must", {"listed-only"}),
        ("nice", {"missing", "listed-only"}),
    )
    selected: list[str] = []
    for priority, evidence_values in tiers:
        for item in skills:
            key = str(item.get("key") or "")
            if (
                key
                and item.get("priority") == priority
                and item.get("evidence") in evidence_values
                and key not in selected
            ):
                selected.append(key)
                if len(selected) >= limit:
                    return selected
    return selected


def rank_projects(parsed: dict[str, Any]) -> list[dict[str, Any]]:
    gaps = [item["name"].lower() for item in parsed["skills"] if item["evidence"] == "missing"]
    family = parsed["job_family"]
    aliases = {"rag": ["rag", "向量数据库", "搜索"], "agent / function calling": ["agent", "function calling", "多智能体"], "llm 评估": ["evals", "模型横评", "测试"], "fastapi": ["fastapi", "api", "工程化"], "docker": ["部署", "工程化"], "sql / 数据库": ["postgresql", "数据工程"]}
    scored = []
    for project in PROJECTS:
        score = 4 if family in project["job_families"] else 0
        score += 2 if project.get("category") == "useful" else 0
        text = " ".join(project["topics"]).lower()
        for gap in gaps:
            terms = aliases.get(gap, [gap])
            score += sum(3 for term in terms if term.lower() in text)
        scored.append((score, project))
    return [project for _, project in sorted(scored, key=lambda item: item[0], reverse=True)]


def fallback_generation(parsed: dict[str, Any], score: int, primary_gap: str, candidates: list[dict[str, Any]]) -> dict[str, Any]:
    asset = parsed["background_assets"][0]
    main = candidates[0]
    primary_gap_key = next(
        (str(item.get("key")) for item in parsed["skills"] if item.get("evidence") == "missing" and item.get("key")),
        str(parsed["skills"][0].get("key") or "skill") if parsed["skills"] else "skill",
    )
    recommendations = []
    for index, project in enumerate(candidates[:3]):
        recommendations.append({
            "project_id": project["id"],
            "reason": f"它能把你的{asset}与「{primary_gap}」缺口交叉起来，并提供可直接复刻的开源实现。",
            "matched_gaps": [primary_gap_key],
            "adaptation": project["copy_angle"],
            "rank": index + 1,
        })
    return {
        "diagnosis": f"你的短板不是关键词不够，而是「{primary_gap}」缺少可验证的项目证据。先补一个能被追问的完整项目，再用固定评测集记录改进前后的差异。",
        "main_project_id": main["id"],
        "project_title": f"基于 {main['name']} 的{asset.replace('背景', '')}证据项目",
        "project_rationale": f"以 {main['full_name']} 为可复制基线，再注入{asset}，避免做成所有人都一样的通用 Demo。",
        "resume_line": f"完成后可填写：基于 {main['name']} 设计并实现面向真实业务样本的 AI 工作流，引入 [样本数量] 条评估集完成错误归因与迭代，将 [核心指标] 从 [基线] 提升至 [结果]。",
        "recommendations": recommendations,
        "milestones": [
            {"week": "01", "title": "复刻最小可运行基线", "deliverable": "跑通仓库示例、记录依赖与输入输出", "talking_point": "为什么选这个仓库作为基线，而不是从零搭建？"},
            {"week": "02", "title": "替换成你的行业数据", "deliverable": f"加入{asset}相关的 20–30 条真实样本", "talking_point": "你的背景怎样改变了任务边界和数据设计？"},
            {"week": "03", "title": "建立评估与错误闭环", "deliverable": "黄金集、失败分类、改进前后对比", "talking_point": "最大失败案例是什么，如何定位并修复？"},
        ],
    }


async def perform_analysis(jd: str, resume: str) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    extraction_system = (ROOT / "evals" / "prompts" / "extractor_v5.txt").read_text(encoding="utf-8").replace(
        "{{ONTOLOGY}}",
        ontology_prompt_block(),
    )
    extraction_user = f"<JD>\n{jd}\n</JD>\n<RESUME>\n{resume}\n</RESUME>"
    extracted, extract_meta = await call_model("extractor", extraction_system, extraction_user, 0)
    parsed = normalize_model_extract(extracted, jd, resume) if extracted else local_extract(jd, resume)
    score, dimensions, primary_gap = score_skills(parsed["skills"])
    candidates = rank_projects(parsed)[:8]
    generation_system = (ROOT / "evals" / "prompts" / "generator_v6.txt").read_text(encoding="utf-8")
    allowed_skill_keys = {str(item["key"]) for item in parsed["skills"]}
    generation_user = json.dumps(
        {
            "score": score,
            "parsed": parsed,
            "background_assets": parsed["background_assets"],
            "allowed_skill_keys": sorted(allowed_skill_keys),
            "required_focus_skill_keys": select_focus_skill_keys(parsed["skills"]),
            "allowed_project_ids": [project["id"] for project in candidates],
            "allowed_projects": candidates,
        },
        ensure_ascii=False,
    )
    generated, generate_meta = await call_model("generator", generation_system, generation_user, 0.3)
    fallback = fallback_generation(parsed, score, primary_gap, candidates)
    if generated:
        generated = normalize_generator_output(generated)
        generator_gate_failures = validate_generator_output(
            generated,
            {project["id"] for project in candidates},
            allowed_skill_keys,
            parsed["background_assets"],
            support_texts=(jd, resume),
        )
        generate_meta["gate_failures"] = generator_gate_failures
        if generator_gate_failures:
            generated = None
    if not generated:
        generated = fallback
    allowed = {project["id"]: project for project in candidates}
    recs = []
    for item in generated.get("recommendations", []):
        project_id = item.get("project_id")
        if project_id not in allowed:
            continue
        recs.append({
            "project": allowed[project_id],
            "reason": str(item.get("reason") or fallback["recommendations"][0]["reason"])[:300],
            "matched_gaps": [str(value)[:40] for value in item.get("matched_gaps", [])[:4]],
            "adaptation": str(item.get("adaptation") or allowed[project_id]["copy_angle"])[:240],
            "rank": len(recs) + 1,
        })
    if len(recs) < 3:
        recs = [{"project": next(project for project in candidates if project["id"] == item["project_id"]), **{key: value for key, value in item.items() if key != "project_id"}} for item in fallback["recommendations"]]
    main_id = generated.get("main_project_id") if generated.get("main_project_id") in allowed else fallback["main_project_id"]
    main_project = allowed[main_id]
    milestones = generated.get("milestones") if isinstance(generated.get("milestones"), list) and len(generated["milestones"]) >= 2 else fallback["milestones"]
    hard_requirement = "暂未发现明显冲突的学历、年限或证书门槛。"
    if re.search(r"硕士|研究生", jd) and not re.search(r"硕士|研究生", resume):
        hard_requirement = "岗位写明硕士优先或要求，当前简历未体现。建议同时搜索本科可投的同类岗位。"
    skills = [
        {
            "key": item["key"],
            "name": item["name"],
            "priority": item["priority"],
            "evidence": item["evidence"],
            "jdQuote": item["jd_quote"],
            "resumeQuote": item["resume_quote"],
            "time": time_for_skill(item["key"]),
        }
        for item in parsed["skills"]
    ]
    result = {
        "role": parsed["role"],
        "jobFamily": parsed["job_family"],
        "score": score,
        "dimensions": dimensions,
        "hardRequirement": hard_requirement,
        "diagnosis": str(generated.get("diagnosis") or fallback["diagnosis"])[:500],
        "backgroundAssets": parsed["background_assets"],
        "skills": skills,
        "project": {
            "title": str(generated.get("project_title") or fallback["project_title"])[:100],
            "rationale": str(generated.get("project_rationale") or fallback["project_rationale"])[:500],
            "duration": main_project["duration"],
            "resumeLine": str(generated.get("resume_line") or fallback["resume_line"])[:600],
            "repository": main_project,
            "milestones": [{"week": str(item.get("week", index + 1)), "title": str(item.get("title", "完成里程碑"))[:80], "deliverable": str(item.get("deliverable", "形成可展示交付物"))[:240], "talkingPoint": str(item.get("talking_point", "说明你的技术选择与权衡"))[:240]} for index, item in enumerate(milestones[:6])],
        },
        "quickWins": [{"title": f"{primary_gap} 最小实验", "duration": "90 分钟", "outcome": "用 5 条样本跑通最小链路"}, {"title": "建立项目评估表", "duration": "半天", "outcome": "定义准确性、可追溯性和失败类型"}, {"title": "写一页技术决策记录", "duration": "1 天", "outcome": "沉淀选型、踩坑与权衡"}],
        "recommendations": recs,
        "source": "model" if extracted or model_config("generator")["api_key"] and generate_meta["schema_ok"] else "rules",
        "model": model_config("generator")["model"] if generate_meta["schema_ok"] else "本地可复现规则",
    }
    return result, [extract_meta, generate_meta]


@app.get("/api/health")
def health() -> dict[str, Any]:
    roles = {
        role: {
            "configured": all(model_config(role)[key] for key in ("base_url", "api_key", "model")),
            "model": model_config(role)["model"] or None,
            "profile": model_config(role)["profile"] or None,
        }
        for role in ("extractor", "generator", "judge_a", "judge_b")
    }
    return {"ok": True, "version": app.version, "models": roles}


@app.post("/api/documents/extract")
async def extract_document(
    request: Request,
    file: UploadFile = File(...),
    kind: str = Query(default="resume", pattern="^(jd|resume)$"),
) -> dict[str, Any]:
    enforce_rate_limit(request, "documents")
    filename = Path(file.filename or "document").name
    content = await file.read(DOCUMENT_MAX_BYTES + 1)
    await file.close()
    if not content:
        raise HTTPException(status_code=422, detail="上传的文件是空的。")
    if len(content) > DOCUMENT_MAX_BYTES:
        raise HTTPException(status_code=413, detail="文件不能超过 10 MB。")
    try:
        raw_text, method, needs_review = _extract_uploaded_document(filename, content)
    except ValueError as error:
        raise HTTPException(status_code=422, detail=str(error)) from error
    except RuntimeError as error:
        raise HTTPException(status_code=503, detail=str(error)) from error
    except Exception as error:
        raise HTTPException(status_code=422, detail="文件无法解析。请确认文件未加密或损坏后重试。") from error

    text = _normalize_document_text(raw_text)
    if len(text) < 20:
        raise HTTPException(status_code=422, detail="没有识别到足够的文字，请上传更清晰的文件或直接粘贴文本。")
    limit = DOCUMENT_TEXT_LIMITS[kind]
    truncated = len(text) > limit
    return {
        "filename": filename,
        "text": text[:limit],
        "characters": min(len(text), limit),
        "method": method,
        "needsReview": needs_review,
        "truncated": truncated,
    }


@app.post("/api/auth/register")
def register(payload: AuthPayload, request: Request) -> dict[str, Any]:
    enforce_rate_limit(request, "register")
    email = payload.email.strip().lower()
    if not re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", email):
        raise HTTPException(status_code=422, detail="请输入有效邮箱。")
    try:
        with db() as connection:
            cursor = connection.execute("INSERT INTO users(email,password_hash,created_at) VALUES(?,?,?)", (email, hash_password(payload.password), now_iso()))
            user_id = cursor.lastrowid
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=409, detail="该邮箱已经注册。")
    return {"token": make_token(user_id, email), "user": {"id": user_id, "email": email}}


@app.post("/api/auth/login")
def login(payload: AuthPayload, request: Request) -> dict[str, Any]:
    enforce_rate_limit(request, "login")
    email = payload.email.strip().lower()
    with db() as connection:
        row = connection.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
    if not row or not verify_password(payload.password, row["password_hash"]):
        raise HTTPException(status_code=401, detail="邮箱或密码不正确。")
    return {"token": make_token(row["id"], row["email"]), "user": {"id": row["id"], "email": row["email"]}}


@app.get("/api/auth/me")
def me(user: dict[str, Any] = Depends(require_user)) -> dict[str, Any]:
    return {"user": {"id": user["sub"], "email": user["email"]}}


@app.delete("/api/account")
def delete_account(request: Request, user: dict[str, Any] = Depends(require_user)) -> dict[str, Any]:
    with db() as connection:
        connection.execute("DELETE FROM model_runs WHERE analysis_id IN (SELECT id FROM analyses WHERE user_id=?)", (user["sub"],))
        connection.execute("DELETE FROM events WHERE user_id=?", (user["sub"],))
        connection.execute("DELETE FROM feedback WHERE user_id=?", (user["sub"],))
        connection.execute("DELETE FROM analyses WHERE user_id=?", (user["sub"],))
        connection.execute("DELETE FROM users WHERE id=?", (user["sub"],))
    return {"ok": True, "requestId": getattr(request.state, "request_id", None)}


@app.post("/api/analyses")
async def create_analysis(payload: AnalysisPayload, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
    user = optional_user(authorization)
    enforce_rate_limit(request, "analysis", user)
    enforce_analysis_budget(request, payload, user)
    result, model_runs = await perform_analysis(payload.jd, payload.resume)
    with db() as connection:
        cursor = connection.execute(
            "INSERT INTO analyses(user_id,role,score,jd_raw,resume_raw,result_json,source,created_at) VALUES(?,?,?,?,?,?,?,?)",
            (user["sub"] if user else None, result["role"], result["score"], payload.jd, payload.resume, json.dumps(result, ensure_ascii=False), result["source"], now_iso()),
        )
        analysis_id = cursor.lastrowid
        for run in model_runs:
            connection.execute("INSERT INTO model_runs(analysis_id,role,model,latency_ms,input_tokens,output_tokens,schema_ok,error,created_at) VALUES(?,?,?,?,?,?,?,?,?)", (analysis_id, run["role"], run["model"], run["latency_ms"], run["input_tokens"], run["output_tokens"], int(run["schema_ok"]), run["error"], now_iso()))
    result["analysisId"] = analysis_id
    result["requestId"] = getattr(request.state, "request_id", None)
    return result


def _prune_analysis_jobs() -> None:
    cutoff = time.time() - ANALYSIS_JOB_TTL_SECONDS
    with _analysis_jobs_lock:
        stale = [job_id for job_id, job in _analysis_jobs.items() if job.get("updated_at", 0) < cutoff]
        for job_id in stale:
            _analysis_jobs.pop(job_id, None)


def _save_analysis_result(
    result: dict[str, Any],
    model_runs: list[dict[str, Any]],
    payload: AnalysisPayload,
    user: dict[str, Any] | None,
) -> int:
    with db() as connection:
        cursor = connection.execute(
            "INSERT INTO analyses(user_id,role,score,jd_raw,resume_raw,result_json,source,created_at) VALUES(?,?,?,?,?,?,?,?)",
            (user["sub"] if user else None, result["role"], result["score"], payload.jd, payload.resume, json.dumps(result, ensure_ascii=False), result["source"], now_iso()),
        )
        analysis_id = cursor.lastrowid
        for run in model_runs:
            connection.execute(
                "INSERT INTO model_runs(analysis_id,role,model,latency_ms,input_tokens,output_tokens,schema_ok,error,created_at) VALUES(?,?,?,?,?,?,?,?,?)",
                (analysis_id, run["role"], run["model"], run["latency_ms"], run["input_tokens"], run["output_tokens"], int(run["schema_ok"]), run["error"], now_iso()),
            )
    return int(analysis_id)


async def _run_analysis_job(job_id: str, payload: AnalysisPayload, user: dict[str, Any] | None) -> None:
    def update(**values: Any) -> None:
        with _analysis_jobs_lock:
            if job_id in _analysis_jobs:
                _analysis_jobs[job_id].update(values, updated_at=time.time())

    update(status="running", stage="analyzing")
    try:
        result, model_runs = await perform_analysis(payload.jd, payload.resume)
        update(stage="saving")
        analysis_id = await asyncio.to_thread(_save_analysis_result, result, model_runs, payload, user)
        result["analysisId"] = analysis_id
        update(status="completed", stage="completed", result=result)
    except Exception as error:
        update(status="failed", stage="failed", error="分析暂时失败，请稍后重试。", error_type=type(error).__name__)


@app.post("/api/analysis-jobs")
async def create_analysis_job(
    payload: AnalysisPayload,
    background_tasks: BackgroundTasks,
    request: Request,
    authorization: str | None = Header(default=None),
) -> dict[str, Any]:
    user = optional_user(authorization)
    enforce_rate_limit(request, "analysis", user)
    enforce_analysis_budget(request, payload, user)
    _prune_analysis_jobs()
    job_id = secrets.token_urlsafe(18)
    with _analysis_jobs_lock:
        _analysis_jobs[job_id] = {
            "status": "queued",
            "stage": "queued",
            "result": None,
            "error": None,
            "owner_id": user["sub"] if user else None,
            "updated_at": time.time(),
        }
    background_tasks.add_task(_run_analysis_job, job_id, payload, user)
    return {"jobId": job_id, "status": "queued", "requestId": getattr(request.state, "request_id", None)}


@app.get("/api/analysis-jobs/{job_id}")
def get_analysis_job(job_id: str, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
    _prune_analysis_jobs()
    user = optional_user(authorization)
    with _analysis_jobs_lock:
        job = _analysis_jobs.get(job_id)
        snapshot = dict(job) if job else None
    if not snapshot or (snapshot.get("owner_id") is not None and (not user or snapshot["owner_id"] != user["sub"])):
        raise HTTPException(status_code=404, detail="没有找到这份分析任务。")
    return {
        "jobId": job_id,
        "status": snapshot["status"],
        "stage": snapshot["stage"],
        "result": snapshot.get("result"),
        "error": snapshot.get("error"),
        "requestId": getattr(request.state, "request_id", None),
    }


def _analysis_visible(analysis_id: int, user: dict[str, Any] | None) -> bool:
    with db() as connection:
        row = connection.execute("SELECT user_id FROM analyses WHERE id=?", (analysis_id,)).fetchone()
    if not row:
        return False
    # Anonymous telemetry may reference an analysis ID, while authenticated
    # users can only annotate their own history.
    return user is None or row["user_id"] in (None, user["sub"])


@app.post("/api/events")
def create_event(payload: EventPayload, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
    user = optional_user(authorization)
    enforce_rate_limit(request, "events", user)
    if payload.analysis_id is not None and not _analysis_visible(payload.analysis_id, user):
        raise HTTPException(status_code=404, detail="找不到关联的分析记录。")
    try:
        metadata_json = json.dumps(payload.metadata, ensure_ascii=False, separators=(",", ":"))
    except (TypeError, ValueError) as error:
        raise HTTPException(status_code=422, detail="事件元数据必须是 JSON 对象。") from error
    if len(metadata_json) > 4096:
        raise HTTPException(status_code=413, detail="事件元数据不能超过 4 KB。")
    with db() as connection:
        cursor = connection.execute(
            "INSERT INTO events(user_id,analysis_id,event,metadata_json,created_at) VALUES(?,?,?,?,?)",
            (user["sub"] if user else None, payload.analysis_id, payload.event, metadata_json, now_iso()),
        )
    return {"ok": True, "eventId": cursor.lastrowid, "requestId": getattr(request.state, "request_id", None)}


@app.post("/api/feedback")
def create_feedback(payload: FeedbackPayload, request: Request, authorization: str | None = Header(default=None)) -> dict[str, Any]:
    user = optional_user(authorization)
    enforce_rate_limit(request, "feedback", user)
    if payload.analysis_id is not None and not _analysis_visible(payload.analysis_id, user):
        raise HTTPException(status_code=404, detail="找不到关联的分析记录。")
    with db() as connection:
        cursor = connection.execute(
            "INSERT INTO feedback(user_id,analysis_id,positive,comment,created_at) VALUES(?,?,?,?,?)",
            (
                user["sub"] if user else None,
                payload.analysis_id,
                None if payload.positive is None else int(payload.positive),
                payload.comment.strip() if payload.comment else None,
                now_iso(),
            ),
        )
    return {"ok": True, "feedbackId": cursor.lastrowid, "requestId": getattr(request.state, "request_id", None)}


@app.get("/api/analyses")
def list_analyses(user: dict[str, Any] = Depends(require_user)) -> list[dict[str, Any]]:
    with db() as connection:
        rows = connection.execute("SELECT id,role,score,source,created_at FROM analyses WHERE user_id=? ORDER BY id DESC", (user["sub"],)).fetchall()
    return [dict(row) for row in rows]


@app.get("/api/analyses/{analysis_id}")
def get_analysis(analysis_id: int, user: dict[str, Any] = Depends(require_user)) -> dict[str, Any]:
    with db() as connection:
        row = connection.execute("SELECT * FROM analyses WHERE id=? AND user_id=?", (analysis_id, user["sub"])).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="没有找到这份分析。")
    result = json.loads(row["result_json"])
    result["analysisId"] = row["id"]
    return result


@app.get("/api/projects")
def projects(
    search: str = Query(default=""),
    job_family: str = Query(default=""),
    category: str = Query(default=""),
    topic: str = Query(default=""),
    difficulty: str = Query(default=""),
    duration: str = Query(default=""),
) -> list[dict[str, Any]]:
    values = PROJECTS
    if search:
        needle = search.lower()
        values = [item for item in values if needle in f"{item['name']} {item['full_name']} {item['description']} {' '.join(item['topics'])}".lower()]
    if job_family:
        values = [item for item in values if job_family in item["job_families"]]
    if category:
        values = [item for item in values if item.get("category") == category]
    if topic:
        values = [item for item in values if topic in item["topics"]]
    if difficulty:
        values = [item for item in values if item["difficulty"] == difficulty]
    if duration:
        values = [item for item in values if item["duration"] == duration]
    return values


@app.get("/api/briefs")
def briefs(window: str = Query(default="24h"), authorization: str | None = Header(default=None)) -> dict[str, Any]:
    user = optional_user(authorization)
    latest = None
    if user:
        with db() as connection:
            row = connection.execute("SELECT result_json FROM analyses WHERE user_id=? ORDER BY id DESC LIMIT 1", (user["sub"],)).fetchone()
        if row:
            latest = json.loads(row["result_json"])
    family = latest.get("jobFamily", "ai_app_dev") if latest else "ai_app_dev"
    project_title = latest.get("project", {}).get("title", "你的主项目") if latest else "你的主项目"
    deep_card = {
        "event": "AI 应用招聘越来越重视评估与失败分析，而不只是能否调用模型。",
        "relationship": f"这与 {family} 岗位直接相关，也能补强你当前项目的工程深度。",
        "question": "你如何证明这个 AI 功能真的变好了，而不是只看几个演示案例？",
        "answer": f"我会先定义任务成功率、引用正确率和失败类型，再用固定样本做回归。以「{project_title}」为例，我保留了改进前后的评估记录。",
        "pitfall": "不要只回答“人工体验不错”，也不要编造不存在的线上指标。",
    }
    return {
        "date": datetime.now().strftime("%Y-%m-%d"),
        "deepCard": deep_card,
        "hotspots": get_daily_hotspots(window=window),
        "items": DAILY_BRIEFS,
    }


FRONTEND_DIST = ROOT / "dist"
if FRONTEND_DIST.is_dir():
    app.mount("/", StaticFiles(directory=FRONTEND_DIST, html=True), name="frontend")
